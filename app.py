import requests
from bs4 import BeautifulSoup
from docx import Document
from urllib.parse import urljoin
import streamlit as st

st.title("Event Scraper to Word Document")

url = st.text_input("Paste the event link here:")

if st.button("Scrape and Generate Word Document"):
    if url:
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')

        doc = Document()

        def add_heading_and_text(doc, heading, text):
            doc.add_heading(heading, level=2)
            doc.add_paragraph(text)

        title_tag = soup.find('h1', class_='event-title')
        title = title_tag.get_text(strip=True) if title_tag else 'N/A'

        date_tag = soup.find('div', class_='event-datetime')
        date = date_tag.get_text(strip=True) if date_tag else 'N/A'

        place_tag = soup.find('div', class_='event-location')
        place = place_tag.get_text(strip=True) if place_tag else 'N/A'

        registration_tag = soup.find('a', class_='event-register-button')
        registration_link = urljoin(url, registration_tag['href']) if registration_tag else 'N/A'

        description_tag = soup.find('div', class_='event-description')
        description = description_tag.get_text(strip=True) if description_tag else 'N/A'

        participants = []
        speakers_section = soup.find('section', class_='event-speakers')
        if speakers_section:
            speaker_names = speakers_section.find_all('div', class_='speaker-name')
            speaker_titles = speakers_section.find_all('div', class_='speaker-title')
            for name, role in zip(speaker_names, speaker_titles):
                participant_entry = f"{name.get_text(strip=True)} ({role.get_text(strip=True)})"
                participants.append(participant_entry)

        add_heading_and_text(doc, 'Event Title:', title)
        add_heading_and_text(doc, 'Date:', date)
        add_heading_and_text(doc, 'Place:', place)
        add_heading_and_text(doc, 'Registration Link:', registration_link)

        if participants:
            add_heading_and_text(doc, 'Participants:', '\\n'.join(participants))
        else:
            add_heading_and_text(doc, 'Participants:', 'No participants listed')

        add_heading_and_text(doc, 'Description:', description)

        doc.save('/tmp/event_info.docx')

        with open('/tmp/event_info.docx', "rb") as file:
            st.download_button(
                label="Download Word Document",
                data=file,
                file_name="event_info.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("Please enter a URL.")
